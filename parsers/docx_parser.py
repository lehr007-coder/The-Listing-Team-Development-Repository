"""
Parse Word (.docx) documents — extracts text content and all embedded images.
"""

import io
import os
from dataclasses import dataclass, field

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image


@dataclass
class ParsedBlog:
    """Container for parsed blog content."""

    title: str = ""
    body_text: str = ""
    paragraphs: list[str] = field(default_factory=list)
    images: list[dict] = field(default_factory=list)  # [{filename, bytes, mime}]

    def full_text(self) -> str:
        """Return the complete text with paragraphs joined."""
        return "\n\n".join(self.paragraphs)


def parse_docx(file_path_or_bytes) -> ParsedBlog:
    """
    Parse a .docx file and extract all text and embedded images.

    Args:
        file_path_or_bytes: Either a file path string or a BytesIO object.

    Returns:
        ParsedBlog with title, paragraphs, and images extracted.
    """
    if isinstance(file_path_or_bytes, (str, os.PathLike)):
        doc = Document(str(file_path_or_bytes))
    else:
        doc = Document(file_path_or_bytes)

    blog = ParsedBlog()

    # --- Extract text ---
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Use the first heading or first paragraph as the title
        if not blog.title:
            style_name = (para.style.name or "").lower()
            if "heading" in style_name or "title" in style_name or i == 0:
                blog.title = text
                continue

        paragraphs.append(text)

    blog.paragraphs = paragraphs
    blog.body_text = "\n\n".join(paragraphs)

    # --- Extract embedded images ---
    image_index = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            content_type = rel.target_part.content_type  # e.g. image/png
            ext = content_type.split("/")[-1] if "/" in content_type else "png"
            # Normalize jpeg
            if ext == "jpeg":
                ext = "jpg"

            image_index += 1
            filename = f"image_{image_index:03d}.{ext}"

            blog.images.append(
                {
                    "filename": filename,
                    "bytes": image_data,
                    "mime": content_type,
                }
            )

    return blog
